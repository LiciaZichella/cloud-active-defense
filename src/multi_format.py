import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import json
import boto3
import os
from datetime import datetime
from urllib.parse import urlparse, parse_qs
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from config import CONFIG

NOME_BUCKET = CONFIG['buckets']['documents']
BUCKET_LOGS = CONFIG['buckets']['audit_logs']

s3 = boto3.client(
    's3',
    endpoint_url=CONFIG['localstack']['endpoint'],
    aws_access_key_id=CONFIG['localstack']['access_key'],
    aws_secret_access_key=CONFIG['localstack']['secret_key'],
    region_name=CONFIG['localstack']['region']
)


def _upload(percorso, nome_file):
    try:
        s3.upload_file(percorso, NOME_BUCKET, nome_file)
        print(f"SI - Documento caricato: {nome_file}")
    except Exception as e:
        print(f"NO - Errore upload: {e}")


def crea_documento_docx(nome_file, titolo, contenuto, beacon_url, prefisso_id, autore):
    print(f"\nGenerazione DOCX: {nome_file}")
    percorso = os.path.join(os.path.dirname(__file__), nome_file)

    doc = Document()
    doc.core_properties.author = autore

    doc.add_heading(titolo, level=1)
    for riga in contenuto.split('\n'):
        doc.add_paragraph(riga)

    if "REAL" in prefisso_id:
        # INCLUDEPICTURE field: Word fetcha l'immagine all'apertura → attiva il beacon
        para = doc.add_paragraph()
        run = para.add_run()

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' INCLUDEPICTURE "{beacon_url}" \\d '
        run._r.append(instr)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)

        print("[*] Iniettato Web Beacon INCLUDEPICTURE nel DOCX.")
    else:
        print("[*] DOCX Honeyfile generato. Allarme configurato sul download.")

    doc.save(percorso)
    print(f"DOCX creato: {nome_file}")
    _upload(percorso, nome_file)
    try:
        beacon_id = parse_qs(urlparse(beacon_url).query).get('file_id', ['SCONOSCIUTO'])[0]
        mapping = {
            "beacon_id": beacon_id,
            "file_name": nome_file,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "tipo": prefisso_id
        }
        s3.put_object(
            Bucket=BUCKET_LOGS,
            Key=f"beacon_mapping/{beacon_id}.json",
            Body=json.dumps(mapping),
            ContentType="application/json"
        )
        print(f"[*] Beacon mapping salvato: {beacon_id} -> {nome_file}")
    except Exception as e:
        print(f"NO - Errore salvataggio beacon mapping: {e}")


def crea_documento_xlsx(nome_file, titolo, contenuto, beacon_url, prefisso_id, autore):
    print(f"\nGenerazione XLSX: {nome_file}")
    percorso = os.path.join(os.path.dirname(__file__), nome_file)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dati"

    if "REAL" in prefisso_id:
        ws['A1'] = f'=HYPERLINK("{beacon_url}", "Apri report")'
        print("[*] Iniettato Web Beacon HYPERLINK nell'XLSX.")
    else:
        ws['A1'] = titolo
        print("[*] XLSX Honeyfile generato. Allarme configurato sul download.")

    for i, riga in enumerate(contenuto.split('\n'), start=3):
        ws[f'A{i}'] = riga

    wb.save(percorso)
    print(f"XLSX creato: {nome_file}")
    _upload(percorso, nome_file)
    try:
        beacon_id = parse_qs(urlparse(beacon_url).query).get('file_id', ['SCONOSCIUTO'])[0]
        mapping = {
            "beacon_id": beacon_id,
            "file_name": nome_file,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "tipo": prefisso_id
        }
        s3.put_object(
            Bucket=BUCKET_LOGS,
            Key=f"beacon_mapping/{beacon_id}.json",
            Body=json.dumps(mapping),
            ContentType="application/json"
        )
        print(f"[*] Beacon mapping salvato: {beacon_id} -> {nome_file}")
    except Exception as e:
        print(f"NO - Errore salvataggio beacon mapping: {e}")
